from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import pandas as pd

# Paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_PATH = os.path.join(BASE_DIR, 'copilot_dev', 'out_lightweight', 'diagnostics_plots', 'pair_account_behavior.png')
DIAG_CSV = os.path.join(BASE_DIR, 'copilot_dev', 'out_lightweight', 'diagnostics.csv')
OUT_DOC = os.path.join(BASE_DIR, 'AnomalyDetection_Detailed.docx')

# Create document
doc = Document()

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

# Title
p = doc.add_paragraph()
r = p.add_run('Anomaly Detection Pipeline - Detailed Documentation')
r.bold = True
r.font.size = Pt(20)
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Executive summary
add_heading('Executive Summary', level=1)
doc.add_paragraph(
    'This document provides a comprehensive description of the anomaly detection pipeline implemented in this repository. '
    'It covers the data schema, modeling approach (Negative Binomial NB2), model selection, change-point handling, diagnostics, and visualization. '
    'An embedded example plot and a slice of diagnostics are included for quick validation.'
)

# Data schema
add_heading('Data Schema and Preparation', level=1)
doc.add_paragraph('Expected input CSV columns:')
cols = ['insight_date_time (datetime)', 'target_type (string)', 'insight_type (string)', 'target_key (string)', 'other columns (optional)']
for c in cols:
    doc.add_paragraph(f'- {c}')

doc.add_paragraph('\nWeekly aggregation: each record is converted to the start-of-week bucket using Monday as week start. Missing weeks are filled with zero counts for each (target_type, insight_type, target_key).')

# Modeling details
add_heading('Modeling Details', level=1)
doc.add_paragraph('We model weekly counts using a Negative Binomial NB2 (mean-variance relationship Var(y)=mu+alpha*mu^2). Key points:')
doc.add_paragraph('- Linear predictor: intercept, optional Fourier terms for seasonality (K harmonics), and a light linear trend (t / 52).')
doc.add_paragraph('- Multiple candidate models are fit per series varying fourier_K; best model is chosen by AIC.')

doc.add_paragraph('Prediction intervals are computed by transforming the model-predicted mean into NB parameters and using the Negative Binomial quantiles (scipy.stats.nbinom.ppf). This yields integer-aware PI appropriate for counts.')

# Change point detection
add_heading('Change Point Detection', level=1)
doc.add_paragraph('We detect structural changes using the ruptures library with the PELT algorithm and an l2 cost on log1p-transformed counts. Steps:')
steps = [
    'Transform counts with log1p for stability.',
    'Run rpt.Pelt(model="l2").fit(transformed_series) and predict with a selected penalty.',
    'If change points are found, split the series and fit models independently on each segment.',
]
for s in steps:
    doc.add_paragraph(f'- {s}')

# Diagnostics
add_heading('Diagnostics & Interpretation', level=1)
doc.add_paragraph('The pipeline outputs a diagnostics CSV with these fields per segment:')
for f in ['target_type','insight_type','target_key (optional)','segment_start','segment_end','aic','fallback (bool)','change_point (start index if any)']:
    doc.add_paragraph(f'- {f}')

doc.add_paragraph('Use these diagnostics to:')
doc.add_paragraph('- Identify series where model fitting failed and fallback was used.')
doc.add_paragraph('- Inspect AIC values to compare model complexity choices per segment.')
doc.add_paragraph('- Locate change points and inspect pre/post behavior.')

# Example visualization
add_heading('Example Visualization (embedded)', level=1)
if os.path.exists(IMG_PATH):
    doc.add_paragraph('The plot below is an example output for the pair-level series "account behavior". It contains model predictions, prediction intervals, anomalies, and change point lines.')
    doc.add_picture(IMG_PATH, width=Inches(6))
    doc.add_paragraph('Figure explanation:')
    doc.add_paragraph('- Blue line: Actual weekly counts')
    doc.add_paragraph('- Dashed line: Model-predicted mean')
    doc.add_paragraph('- Gray band: Prediction interval (NB quantiles)')
    doc.add_paragraph('- Red dots: Points flagged as anomalies (outside PI)')
    doc.add_paragraph('- Purple vertical lines: Detected change points where the model was refit')
else:
    doc.add_paragraph('Example plot not found at: ' + IMG_PATH)

# Embed a small diagnostics table if available
add_heading('Diagnostics Sample (first 10 rows)', level=1)
if os.path.exists(DIAG_CSV):
    try:
        df = pd.read_csv(DIAG_CSV)
        df_small = df.head(10)
        table = doc.add_table(rows=1, cols=len(df_small.columns))
        hdr_cells = table.rows[0].cells
        for i, c in enumerate(df_small.columns):
            hdr_cells[i].text = str(c)
        for _, row in df_small.iterrows():
            cells = table.add_row().cells
            for i, c in enumerate(df_small.columns):
                cells[i].text = str(row[c])
    except Exception as e:
        doc.add_paragraph('Failed to read diagnostics CSV: ' + str(e))
else:
    doc.add_paragraph('Diagnostics CSV not found at: ' + DIAG_CSV)

# Implementation notes and reproducibility
add_heading('Implementation Notes & Reproducibility', level=1)
doc.add_paragraph('Key files and where to find them:')
files = [
    'copilot_dev/lightweight_nb_anomaly.py - main pipeline',
    'copilot_dev/run_lightweight.py - runner script',
    'copilot_dev/requirements.txt - dependencies for enhanced pipeline',
]
for f in files:
    doc.add_paragraph('- ' + f)

doc.add_paragraph('\nReproducibility checklist:')
doc.add_paragraph('- Ensure virtualenv is activated and dependencies from requirements.txt are installed.')
doc.add_paragraph('- Run the runner script to regenerate plots and diagnostics.')

# Code snippets
add_heading('Key Code Snippets', level=1)
code_snippet = '''# build_design: create design matrix with Fourier terms
P = 52.0
X = [np.ones(len(g))]
for k in range(1, K+1):
    X.append(np.sin(2*np.pi*k*g['t']/P))
    X.append(np.cos(2*np.pi*k*g['t']/P))
X.append(g['t']/P)
X = np.column_stack(X)
'''
pre = doc.add_paragraph()
pre.add_run(code_snippet).font.name = 'Courier New'

# Gen AI prompt
add_heading('Gen AI Prompt (for Copilot/ChatGPT)', level=1)
doc.add_paragraph('Use the prompt below to reproduce or adapt this pipeline for a new dataset:')
doc.add_paragraph(
    "I have weekly count data in CSV format, with columns like `target_type`, `insight_type`, `target_key`, `insight_date_time`, and I want to detect anomalies for each (target_type, insight_type) pair and key.\n"
    "Please develop a robust Python pipeline that:\n"
    "- Aggregates the data to weekly counts for each (target_type, insight_type, target_key).\n"
    "- For each (target_type, insight_type) pair and key:\n"
    "    - Fits multiple Negative Binomial (NB2) models with different numbers of Fourier terms (for seasonality) and selects the best model using AIC.\n"
    "    - Detects change points in the time series using the `ruptures` library (with a log1p transform and `model='l2'`), and if a change point is found, fits separate models to each segment.\n"
    "    - Uses a robust rolling window fallback if model fitting fails.\n"
    "    - Calculates prediction intervals using the NB2 model and flags anomalies outside the interval.\n"
    "    - Collects diagnostics (AIC, fallback usage, change points) for each segment and saves them to a CSV.\n"
    "    - Generates and saves plots for each series, showing actuals, predictions, intervals, anomalies, and change points.\n"
    "- Outputs:\n"
    "    - Anomaly flags for each key and pair as CSVs.\n"
    "    - Diagnostics CSV.\n"
    "    - Plots for each series in a diagnostics folder.\n"
    "- Use only open-source Python libraries (pandas, numpy, statsmodels, ruptures, matplotlib).\n"
    "- The code should be robust, modular, and ready to run on new data with similar structure.\n\n"
    "Please provide the full code, including requirements, and ensure it is thoroughly tested and handles edge cases."
)

# Save document
os.makedirs(BASE_DIR, exist_ok=True)
doc.save(OUT_DOC)
print('Saved detailed documentation to:', OUT_DOC)

